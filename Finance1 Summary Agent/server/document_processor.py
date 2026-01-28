"""
Document processing utilities for file uploads.
Handles text extraction and chunking for vectorstore ingestion.
"""
from typing import List
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter


def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extracts text content from uploaded file.
    
    Args:
        file_content: Raw file bytes
        filename: Original filename (used to determine file type)
    
    Returns:
        Extracted text content
    """
    # Get file extension
    file_ext = filename.split(".")[-1].lower() if "." in filename else ""
    
    # Handle text files
    if file_ext in ["txt", "text", "md", "markdown"]:
        try:
            return file_content.decode("utf-8", errors="ignore")
        except Exception as e:
            raise ValueError(f"Error decoding text file: {e}")
    
    # Handle PDF files (requires PyPDF2 or pdfplumber)
    elif file_ext == "pdf":
        try:
            import PyPDF2
            from io import BytesIO
            
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_parts = []
            
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())
            
            return "\n".join(text_parts)
        except ImportError:
            raise ValueError("PDF processing requires PyPDF2. Install with: pip install PyPDF2")
        except Exception as e:
            raise ValueError(f"Error processing PDF: {e}")
    
    # Handle DOCX files (requires python-docx)
    elif file_ext in ["docx", "doc"]:
        try:
            from docx import Document as DocxDocument
            from io import BytesIO
            
            docx_file = BytesIO(file_content)
            doc = DocxDocument(docx_file)
            text_parts = [paragraph.text for paragraph in doc.paragraphs]
            
            return "\n".join(text_parts)
        except ImportError:
            raise ValueError("DOCX processing requires python-docx. Install with: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Error processing DOCX: {e}")
    
    # Default: try to decode as text
    else:
        try:
            return file_content.decode("utf-8", errors="ignore")
        except Exception:
            raise ValueError(f"Unsupported file type: {file_ext}")


def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
    """
    Splits text into chunks for vectorstore ingestion.
    
    Args:
        text: Text content to chunk
        chunk_size: Maximum size of each chunk
        chunk_overlap: Overlap between chunks
    
    Returns:
        List of text chunks
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
    )
    
    chunks = text_splitter.split_text(text)
    return chunks


def process_document(file_content: bytes, filename: str, entity_id: str, entity_type: str) -> List[Document]:
    """
    Processes a document file and returns Document objects ready for vectorstore.
    
    Args:
        file_content: Raw file bytes
        filename: Original filename
        entity_id: Associated entity ID
        entity_type: Associated entity type
    
    Returns:
        List of Document objects with metadata
    """
    # Extract text
    text = extract_text_from_file(file_content, filename)
    
    # Chunk text
    chunks = chunk_text(text)
    
    # Create Document objects with metadata
    documents = [
        Document(
            page_content=chunk,
            metadata={
                "source": filename,
                "entity_id": entity_id,
                "entity_type": entity_type,
                "chunk_index": i
            }
        )
        for i, chunk in enumerate(chunks)
    ]
    
    return documents
